import os
import re
from typing import List, Tuple

from pypdf import PdfReader  # ✅ pypdf is the maintained fork of PyPDF2
from docx import Document


def _text_quality_score(s: str) -> int:
    """Favor extractions with real words over noise or empty layout artifacts."""
    if not s:
        return 0
    letters = sum(1 for c in s if c.isalpha())
    words = len(re.findall(r"[A-Za-z]{2,}", s))
    return letters + 3 * words


def _extract_pdf_pymupdf(file_path: str) -> str:
    """PyMuPDF often succeeds where pypdf returns little (multi-column, odd encodings)."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        try:
            parts: List[str] = []
            for page in doc:
                t = page.get_text("text")
                if t:
                    parts.append(t)
            return "\n".join(parts).strip()
        finally:
            doc.close()
    except Exception:
        return ""


def _extract_pdf_pypdf(file_path: str) -> str:
    try:
        reader = PdfReader(file_path)
        parts: List[str] = []
        for page in reader.pages:
            t = page.extract_text()
            if not t or len(t.strip()) < 30:
                try:
                    t2 = page.extract_text(extraction_mode="layout")
                    if t2 and len(t2.strip()) > len((t or "").strip()):
                        t = t2
                except TypeError:
                    pass
            if t:
                parts.append(t)
        return "\n".join(parts).strip()
    except Exception:
        return ""


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF using multiple engines and keep the best result.
    Image-only (scanned) PDFs still yield no text without OCR.
    """
    candidates: List[Tuple[str, str]] = []

    mu = _extract_pdf_pymupdf(file_path)
    if mu:
        candidates.append(("pymupdf", mu))

    py = _extract_pdf_pypdf(file_path)
    if py:
        candidates.append(("pypdf", py))

    if not candidates:
        raise ValueError(
            "Could not read PDF (encrypted, corrupt, or image-only). "
            "Try exporting as PDF from Word, or upload .docx / .txt."
        )

    best_label, best_text = max(candidates, key=lambda x: _text_quality_score(x[1]))
    # If both are weak, still return the best attempt for downstream heuristics
    if _text_quality_score(best_text) < 40:
        print(f"⚠️ PDF text is very short ({len(best_text)} chars, via {best_label}) — may be scanned.")
    return best_text


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a .docx Word file (paragraphs and table cells)."""
    try:
        doc = Document(file_path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        parts.append(t)
        return "\n".join(parts)
    except Exception as e:
        raise ValueError(f"Failed to read DOCX: {e}")


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from a plain .txt file."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except UnicodeDecodeError:
        # ✅ Fallback encoding for non-UTF-8 files
        with open(file_path, "r", encoding="latin-1") as f:
            return f.read().strip()
    except Exception as e:
        raise ValueError(f"Failed to read TXT: {e}")


def extract_text(file_path: str) -> str:
    """
    Extract text from a resume file.
    Supports: .pdf, .docx, .txt
    """

    # ✅ Check file exists before processing
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    # ✅ Check file is not empty
    if os.path.getsize(file_path) == 0:
        raise ValueError(f"File is empty: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: '{ext}'. Supported: .pdf, .docx, .txt")